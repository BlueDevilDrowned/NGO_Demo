from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


ROOT = Path(r"D:\UnityHub\Project\NGO")
SOURCE = Path(r"C:\Users\BlueDevil\Downloads\轮回游戏_30条道具线索_NPC条件对话触发表.docx")
DATA_SCRIPT = ROOT / "tools" / "build_complete_dialogue_doc.py"
GRAPH_DIR = ROOT / "generated_story_graphs"
OUTPUT = ROOT / "轮回游戏_角色故事图_可视化对白版.docx"


GRAPH_ORDER = [
    ("LiuChouStoryGraph", "刘丑", "01-liuchou-story-graph.png"),
    ("StudentGroupStoryGraph", "周围学生", "02-student-group-story-graph.png"),
    ("ZhouYaoStoryGraph", "周曜", "03-zhouyao-story-graph.png"),
    ("GuTeacherStoryGraph", "顾老师", "04-gu-teacher-story-graph.png"),
    ("HeTeacherStoryGraph", "贺老师", "05-he-teacher-story-graph.png"),
    ("SuHeStoryGraph", "苏禾", "06-suhe-story-graph.png"),
    ("ChenUncleStoryGraph", "陈叔", "07-chen-uncle-story-graph.png"),
    ("XuChengStoryGraph", "许澄", "08-xucheng-story-graph.png"),
    ("LuoTeacherStoryGraph", "罗老师", "09-luo-teacher-story-graph.png"),
]


GRAPH_CAPTIONS = {
    "LiuChouStoryGraph": "A03在首次普通入口中主动给予；A01取得后才成为可选话题。后续由A05、K02、A15、K04、K05逐步把刘丑从引导者推向防御与最终对质。",
    "StudentGroupStoryGraph": "学生群体只有一张简单图：普通回避，取得B01后开放匿名纸条话题，回答仍保持群体噤声边界。",
    "ZhouYaoStoryGraph": "周曜沿走廊证据推进：拖拽路线、值日表冲突、扶手纤维、校徽背扣。跨角色结果只写入状态，再由刘丑图读取。",
    "GuTeacherStoryGraph": "顾老师的三条线按证据强度推进：长期伤情、旧检修正常、金属碎屑工具损伤。",
    "HeTeacherStoryGraph": "贺老师负责资料与恢复：楼层手绘图、盲区投诉、录音中第三人。",
    "SuHeStoryGraph": "苏禾先谨慎回避；速写与统计纸提供背景，玩家取得至少一项霸凌证据后，她才主动交出A07录音。",
    "ChenUncleStoryGraph": "陈叔的记录线由缺页误导开始，经交接便签反证外来者，再用远景监控与校准记录闭合时间可信度。",
    "XuChengStoryGraph": "许澄负责把零散物证变成可靠关联：名单、封袋、总账、工具残留与伪造申请核验。",
    "LuoTeacherStoryGraph": "罗老师图分为道德责任与办公室设备两支：投诉信确认压案，打印缓存只确认打印地点。",
}


def load_data():
    spec = spec_from_file_location("dialogue_data", DATA_SCRIPT)
    module = module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def add_header_footer(data, section, header_text, footer_text):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    data.set_font(hp.add_run(header_text), size=8.5, color=(100, 111, 118))
    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data.set_font(fp.add_run(footer_text), size=8.5, color=(118, 124, 128))


def add_fitted_picture(doc, image_path, max_width_cm, max_height_cm):
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    scale = min(max_width_cm / width_px, max_height_cm / height_px)
    width_cm = width_px * scale
    height_cm = height_px * scale
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Cm(0.15)
    p.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))


def add_option(data, doc, number, label, lines):
    p = doc.add_paragraph(style="Dialogue Option")
    data.set_font(p.add_run(f"选择 {number}｜{label}"), size=10.5, bold=True, color=(42, 93, 83))
    data.add_dialogue(doc, lines)


def build():
    data = load_data()
    doc = Document(SOURCE)
    data.configure_styles(doc)

    graph_section = doc.add_section(WD_SECTION.NEW_PAGE)
    graph_section.orientation = WD_ORIENT.LANDSCAPE
    graph_section.page_width = Cm(29.7)
    graph_section.page_height = Cm(21.0)
    graph_section.top_margin = Cm(1.0)
    graph_section.bottom_margin = Cm(1.0)
    graph_section.left_margin = Cm(1.2)
    graph_section.right_margin = Cm(1.2)
    graph_section.header_distance = Cm(0.5)
    graph_section.footer_distance = Cm(0.5)
    add_header_footer(data, graph_section, "轮回游戏｜角色故事图", "每个角色一张图｜横向轮次天次，纵向线索分支")

    for index, (graph_name, role, filename) in enumerate(GRAPH_ORDER):
        if index > 0:
            doc.add_page_break()
        p = doc.add_paragraph(style="Script Title")
        p.paragraph_format.space_after = Cm(0.08)
        data.set_font(p.add_run(f"七-{index + 1}｜{role}角色故事图"), size=20, bold=True, color=(31, 67, 89))
        data.add_bottom_border(p, color="B7C9D6", size="5")
        add_fitted_picture(doc, GRAPH_DIR / filename, 26.8, 15.4)
        data.add_note(doc, "图意", GRAPH_CAPTIONS[graph_name], "EEF3F5")

    dialogue_section = doc.add_section(WD_SECTION.NEW_PAGE)
    dialogue_section.orientation = WD_ORIENT.PORTRAIT
    dialogue_section.page_width = Cm(21.0)
    dialogue_section.page_height = Cm(29.7)
    dialogue_section.top_margin = Cm(1.8)
    dialogue_section.bottom_margin = Cm(1.8)
    dialogue_section.left_margin = Cm(2.1)
    dialogue_section.right_margin = Cm(2.1)
    dialogue_section.header_distance = Cm(0.8)
    dialogue_section.footer_distance = Cm(0.8)
    add_header_footer(data, dialogue_section, "轮回游戏｜故事图节点完整对白", "普通入口唯一｜线索取得后开放话题或主动事件")

    title = doc.add_paragraph(style="Script Title")
    data.set_font(title.add_run("八、故事图节点完整对白"), size=22, bold=True, color=(31, 67, 89))
    data.add_bottom_border(title)
    data.add_note(doc, "阅读方式", "先在前面的角色故事图中看流向；需要具体台词时，再按角色、轮次和线索ID查本章。这里不再为每条线索重复普通开场。", "E8F0F2")

    clue_by_id = {clue["id"]: clue for clue in data.CLUES}
    for graph_index, (graph_name, role, _) in enumerate(GRAPH_ORDER):
        group_heading = doc.add_paragraph(style="Clue Heading")
        if graph_index > 0:
            group_heading.paragraph_format.page_break_before = True
        data.set_font(group_heading.add_run(f"{role}｜普通入口与线索节点"), size=16, bold=True, color=(31, 67, 89))
        data.add_bottom_border(group_heading, color="D6E1E6", size="4")

        data.add_heading(doc, "普通会话入口")
        bases = [base for base in data.BASE_CONVERSATIONS if role in base["key"]]
        for base in bases:
            data.add_heading(doc, base["key"])
            data.add_dialogue(doc, base["lines"])
            data.add_note(doc, "可选动作", base["special"], "EDF5F0")
            data.add_note(doc, "重复交谈", base["repeat"], "F3F4F5")

        primary_clues = []
        for clue in data.CLUES:
            assignment, _ = data.GRAPH_ASSIGNMENTS[clue["id"]]
            if assignment.split(" + ", 1)[0] == graph_name:
                primary_clues.append(clue)

        data.add_heading(doc, "线索节点对白")
        for clue_index, clue in enumerate(primary_clues):
            h = doc.add_paragraph(style="Clue Heading")
            if clue_index > 0:
                h.paragraph_format.page_break_before = True
            data.set_font(h.add_run(f"{clue['id']}｜{clue['item']}"), size=15, bold=True, color=(31, 67, 89))
            data.add_bottom_border(h, color="D6E1E6", size="4")

            assignment, node_id = data.GRAPH_ASSIGNMENTS[clue["id"]]
            mode, condition, topic, activation_note = data.ACTIVATION[clue["id"]]
            data.add_note(doc, "位置", f"{clue['round']} · {clue['scene']} · {assignment}", "EEF3F5")
            data.add_note(doc, "进入", f"{mode}｜{topic}｜条件：{condition}", "EDF5F0")
            data.add_note(doc, "说明", activation_note, "F3F4F5")

            data.add_heading(doc, "进入节点后的对白")
            data.add_dialogue(doc, clue["trigger"])
            data.add_heading(doc, "主角选择与后续")
            for option_no, (label, lines) in enumerate(clue["choices"], start=1):
                add_option(data, doc, option_no, label, lines)
            data.add_heading(doc, "联动与收束")
            data.add_note(doc, "联动", clue["link"], "EDF5F0")
            data.add_note(doc, "结果", clue["close"], "F3F4F5")
            data.add_note(doc, "事实边界", clue["boundary"], "F8F1E6")

    doc.core_properties.title = "轮回游戏｜角色故事图可视化对白版"
    doc.core_properties.subject = "每个角色一张故事图，并附普通入口、线索节点与主角分支完整对白"
    doc.core_properties.comments = "保留原始线索条件表；新增九张角色故事图和按图整理的对白脚本。"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
