from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\BlueDevil\Downloads\轮回游戏_30条道具线索_NPC条件对话触发表.docx")
OUTPUT = Path(r"D:\UnityHub\Project\NGO\轮回游戏_角色故事图_可视化对白版.docx")
EXPECTED_IDS = [
    "A01", "A03", "B01", "A04", "A08", "A11", "B06", "A06", "B02", "A12",
    "A02", "A10", "K01", "A05", "K03", "B04", "A07", "B03", "B07", "A13",
    "A14", "A09", "B05", "K02", "B08", "B09", "A15", "K04", "B10", "K05",
]


def table_text(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


def main():
    source = Document(SOURCE)
    output = Document(OUTPUT)
    issues = []

    ids = [
        p.text.split("｜", 1)[0]
        for p in output.paragraphs
        if p.style.name == "Clue Heading" and p.text.split("｜", 1)[0] in EXPECTED_IDS
    ]
    options = [p for p in output.paragraphs if p.style.name == "Dialogue Option"]
    branch_headings = [p.text for p in output.paragraphs if p.style.name == "Branch Heading"]

    if len(ids) != len(EXPECTED_IDS) or set(ids) != set(EXPECTED_IDS):
        issues.append(f"clue ids mismatch: {ids}")
    if len(options) != 90:
        issues.append(f"expected 90 clue options, got {len(options)}")
    if len(output.sections) != 3:
        issues.append(f"expected 3 sections, got {len(output.sections)}")
    if len(output.tables) != 7:
        issues.append(f"expected only the 7 original tables, got {len(output.tables)}")
    if len(output.inline_shapes) != 9:
        issues.append(f"expected 9 character graph images, got {len(output.inline_shapes)}")

    for index, source_table in enumerate(source.tables):
        if table_text(source_table) != table_text(output.tables[index]):
            issues.append(f"source table {index + 1} changed")

    all_text = "\n".join(p.text for p in output.paragraphs)
    required_text = [
        "七-1｜刘丑角色故事图",
        "七-9｜罗老师角色故事图",
        "八、故事图节点完整对白",
        "刘丑｜普通入口与线索节点",
        "罗老师｜普通入口与线索节点",
    ]
    for required in required_text:
        if required not in all_text:
            issues.append(f"required story-graph section missing: {required}")
    for obsolete in ("7.1 故事图求值顺序", "建议节点字段", "九张角色故事图分配"):
        if obsolete in all_text:
            issues.append(f"obsolete program-table content still exists: {obsolete}")
    for placeholder in ("TODO", "待补", "{{", "}}"):
        if placeholder in all_text:
            issues.append(f"placeholder found: {placeholder}")

    with ZipFile(OUTPUT) as archive:
        bad_member = archive.testzip()
    if bad_member:
        issues.append(f"damaged ZIP member: {bad_member}")

    print({
        "issues": len(issues),
        "clues": len(ids),
        "options": len(options),
        "branch_headings": len(branch_headings),
        "tables": len(output.tables),
        "graph_images": len(output.inline_shapes),
        "sections": len(output.sections),
        "bytes": OUTPUT.stat().st_size,
    })
    if issues:
        for issue in issues:
            print("ISSUE:", issue)
        raise SystemExit(1)
    print("OK: 9 graphs, 30 clue nodes, 90 choices, source preservation, and package integrity passed")


if __name__ == "__main__":
    main()
